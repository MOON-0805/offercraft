import io
import json
from fastapi import APIRouter, Depends, UploadFile, File, HTTPException
from sqlalchemy.orm import Session
from backend.database import get_db
from backend.db_sync import schedule_db_upload
from backend import models
from backend.schemas import ResumeCreate, ResumeUpdate, ResumeOut
from backend.ai_client import call_deepseek, build_prompt

router = APIRouter(prefix="/api/resumes", tags=["resumes"])


def _extract_text_from_docx(content: bytes) -> str:
    """Extract plain text from a .docx file using python-docx."""
    from docx import Document
    doc = Document(io.BytesIO(content))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_text_from_pdf(content: bytes) -> str:
    """Extract plain text from a .pdf file using pdfplumber."""
    import pdfplumber
    parts: list[str] = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text and text.strip():
                parts.append(text.strip())
    return "\n".join(parts)


@router.get("", response_model=list[ResumeOut])
def list_resumes(db: Session = Depends(get_db)):
    return db.query(models.Resume).order_by(models.Resume.updated_at.desc()).all()


@router.get("/{resume_id}", response_model=ResumeOut)
def get_resume(resume_id: int, db: Session = Depends(get_db)):
    r = db.query(models.Resume).filter(models.Resume.id == resume_id).first()
    if not r:
        raise HTTPException(404, "简历不存在")
    return r


@router.post("", response_model=ResumeOut)
def create_resume(body: ResumeCreate, db: Session = Depends(get_db)):
    if body.is_default:
        db.query(models.Resume).filter(models.Resume.is_default == True).update({"is_default": False})
    r = models.Resume(**body.model_dump())
    db.add(r)
    db.commit()
    db.refresh(r)
    schedule_db_upload()
    return r


@router.put("/{resume_id}", response_model=ResumeOut)
def update_resume(resume_id: int, body: ResumeUpdate, db: Session = Depends(get_db)):
    r = db.query(models.Resume).filter(models.Resume.id == resume_id).first()
    if not r:
        raise HTTPException(404, "简历不存在")
    data = body.model_dump(exclude_unset=True)
    if data.get("is_default"):
        db.query(models.Resume).filter(models.Resume.is_default == True).update({"is_default": False})
    for k, v in data.items():
        setattr(r, k, v)
    db.commit()
    db.refresh(r)
    schedule_db_upload()
    return r


@router.delete("/{resume_id}")
def delete_resume(resume_id: int, db: Session = Depends(get_db)):
    r = db.query(models.Resume).filter(models.Resume.id == resume_id).first()
    if not r:
        raise HTTPException(404, "简历不存在")
    db.delete(r)
    db.commit()
    schedule_db_upload()
    return {"ok": True}


@router.post("/upload")
async def upload_resume(file: UploadFile = File(...), db: Session = Depends(get_db)):
    """Upload a resume file (.pdf, .docx, .md, .txt) and extract text."""
    filename = file.filename or "resume.txt"
    content = await file.read()
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else "txt"

    text = ""
    source_type = ext

    if ext == "docx":
        try:
            text = _extract_text_from_docx(content)
        except Exception as e:
            raise HTTPException(400, f"无法解析 DOCX 文件: {e}")
    elif ext == "pdf":
        try:
            text = _extract_text_from_pdf(content)
        except Exception as e:
            raise HTTPException(400, f"无法解析 PDF 文件: {e}")
    elif ext in ("md", "markdown"):
        try:
            text = content.decode("utf-8")
        except UnicodeDecodeError:
            text = content.decode("gbk", errors="replace")
        source_type = "md"
    else:  # .txt and others
        try:
            text = content.decode("utf-8")
        except UnicodeDecodeError:
            text = content.decode("gbk", errors="replace")
        source_type = "txt"

    if not text.strip():
        raise HTTPException(400, "无法从文件中提取文本内容，请确认文件非空")

    r = models.Resume(
        name=filename.rsplit(".", 1)[0],
        raw_text=text,
        source_type=source_type,
    )
    db.add(r)
    db.commit()
    db.refresh(r)
    schedule_db_upload()
    return {"id": r.id, "name": r.name, "raw_text": r.raw_text, "source_type": r.source_type}


@router.post("/{resume_id}/parse")
def parse_resume(resume_id: int, db: Session = Depends(get_db)):
    """Use AI to parse raw resume text into structured JSON."""
    r = db.query(models.Resume).filter(models.Resume.id == resume_id).first()
    if not r:
        raise HTTPException(404, "简历不存在")
    if not r.raw_text.strip():
        raise HTTPException(400, "简历内容为空，无法解析")

    from backend.ai_client import call_deepseek, build_prompt
    system = """你是一个简历解析助手。将用户提供的简历文本解析为结构化 JSON。
输出格式严格如下（不要多余文字）：
{
  "basic_info": {"name":"","email":"","phone":"","github":"","blog":""},
  "education": [{"school":"","degree":"","major":"","start":"","end":"","highlights":[]}],
  "experience": [{"company":"","title":"","start":"","end":"","bullets":[]}],
  "projects": [{"name":"","role":"","start":"","end":"","bullets":[],"tech_stack":[]}],
  "skills": ["技能1","技能2"]
}"""
    messages = build_prompt(system, r.raw_text)
    try:
        result = call_deepseek(messages, temperature=0.1, max_tokens=4096)
        # Try to extract JSON from the response
        result = result.strip()
        if result.startswith("```"):
            result = result.split("\n", 1)[-1].rsplit("```", 1)[0].strip()
        parsed = json.loads(result)
        r.basic_info = parsed.get("basic_info", {})
        r.education = parsed.get("education", [])
        r.experience = parsed.get("experience", [])
        r.projects = parsed.get("projects", [])
        r.skills = parsed.get("skills", [])
        db.commit()
        db.refresh(r)
        schedule_db_upload()
        return {"ok": True, "parsed": parsed}
    except json.JSONDecodeError:
        raise HTTPException(500, "AI 返回格式异常，请重试")
    except ValueError as e:
        raise HTTPException(400, str(e))
